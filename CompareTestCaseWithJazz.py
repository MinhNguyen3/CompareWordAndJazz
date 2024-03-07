import docx
import os
import sys
import pandas as pd
from os.path import dirname

working_folder = dirname(__file__)
sys.path.insert(0, working_folder)  

# def compare_applicable (word_applicable, jazz_applicable):
#     # Get applicable products name from word
#     word_applicable = word_applicable.replace(" ","")
#     word_applicable_list = word_applicable.split(",")
#     while "" in word_applicable_list:
#         word_applicable_list.remove("")

#     # Get applicable products name from jazz
#     jazz_applicable = jazz_applicable.replace(" ","")
#     jazz_applicable_list = jazz_applicable.split(",")
#     while "" in jazz_applicable_list:
#         jazz_applicable_list.remove("")

#     # Get the difference
#     difference_applicable = list(set(word_applicable_list) ^ set(jazz_applicable_list))
#     # return_str = ""
#     # for x in difference_applicable:
#     #     return_str += x + ","
#     # if return_str:
#     #     return_str = return_str[:-1]
#     # return return_str
#     return difference_applicable

def analyze_applicable_products_to_list(applicable_products):
    # Get applicable products name from word
    applicable_products = applicable_products.replace(" ","")
    applicable_products_list = applicable_products.split(",")
    while "" in applicable_products_list:
        applicable_products_list.remove("")
    return applicable_products_list

#================================ANALYZE WORD FILE=============================
# document = docx.Document(working_folder + "\\Interface_X28X9.docx")
# document = docx.Document(working_folder + "\\Bluetooth Radio Interface_X1X3.docx")
# document = docx.Document(working_folder + "\\Service Port_X9X15.docx")
# document = docx.Document(working_folder + "\\User Interface_X17X18.docx")
# document = docx.Document(working_folder + "\\Retail-Upgrade Firmware_X1X6.docx")
# document = docx.Document(working_folder + "\\Service Port_X9X15.docx")

PRIORITYTC = {"O" : "Often", "S" : "Sometimes", "R" : "Rare", "V" : "Very often"}
AUTOMANUAL = {"A" : "Automated completely", "P" : "Partially automated", "M" : "Manual Test", "N" : "No implemented"}
TCList = []
TestCase = ""

for file_name in os.listdir(working_folder):
    if file_name.endswith(".docx"):
        print("Reading " + file_name)
        document = docx.Document(working_folder + "\\" + file_name)

        for paragraph in document.paragraphs:
            # Get all paragraph have level 2
            if paragraph.style.name == 'Heading 2':
                text = paragraph.text
                # Check format MT_... (...)[...]
                if ("MT" in text) and ("]" in text):
                    # Get TC Name
                    name = text.split("(")[0].replace(" ","")
                    print(name)

                    # Get Attribute          
                    find_open_round_bracket = text.find("(")
                    find_close_round_bracket = text.find(")")
                    attributes = text[find_open_round_bracket + 1 : find_close_round_bracket].replace(" ","")
                    # print(attributes)

                    attributes_list = attributes.split(attributes[-2])
                    # print(attributes_list)

                    estimate = attributes_list[0]
                    try:
                        estimate = int(estimate)
                    except:
                        estimate = int(estimate[:-1])
                        
                    priority = attributes_list[1]
                    automanual = attributes_list[2]

                    # print(estimate)
                    # print(often)
                    # print(automanual)

                    # Get Applicable Product
                    find_open_square_bracket = text.find("[")
                    find_close_square_bracket = text.find("]")
                    applicable_products = text[find_open_square_bracket + 1 : find_close_square_bracket]
                    applicable_products = applicable_products.replace(" ","")
                    # print(applicable)


                    TC = {"Name" : name,
                        "Estimate" : estimate,
                        "Priority TC" : PRIORITYTC[priority],
                        "Auto/Manual" : AUTOMANUAL[automanual],
                        "Applicable" : applicable_products}
                    
                    TCList.append(TC)
                    
word_tcs = pd.DataFrame(TCList)
# print(word_tcs)
word_tcs.to_excel(working_folder + "\\TestCasesFromWord.xlsx")



#================================ANALYZE JAZZ FILE=============================
jazz_tcs = pd.read_excel(os.getcwd() + "\\TestCasesFromJazz.xlsx")
print("Reading TestCasesFromJazz.xlsx")
jazz_tcs_list = []
# Modify Estimate time to minutes
for x in jazz_tcs["Name"]:
    # print(x)
    # jazz_tcs_list.append(x)
    jazz_tcs_list.append(x.replace(" ",""))
    jazz_estimate = jazz_tcs[jazz_tcs["Name"] == x]["Estimate"].tolist()[0].strip()
    jazz_estimate = jazz_estimate.replace(" min","").replace(" hr ","*60+").replace(" hr","*60")
    jazz_tcs.loc[jazz_tcs["Name"] == x,["Estimate"]] = eval(jazz_estimate)

jazz_tcs["Name"] = jazz_tcs_list
# for x in jazz_tcs["Name"]:
#     if " " in x:
#         print(x)


#===============================COMPARE WORD AND JAZZ==========================
print("Comparing")
# Find TCs are not updated to Jazz
tcs_not_on_jazz = []
for x in word_tcs["Name"]:
    if x not in jazz_tcs_list:
        tcs_not_on_jazz.append(x)
new_tcs = word_tcs[word_tcs["Name"].isin(tcs_not_on_jazz)]

# Compare TCs on Jazz with TCs in Word
df = pd.DataFrame(columns = ["Name", 
                            "Word Estimate", "Jazz Estimate",
                            "Word Priority TC", "Jazz Priority TC",
                            "Word Auto/Manual", "Jazz Auto/Manual",
                            "Applicable"])


Difference_TCs = []
TheSame_TCs = []
for x in word_tcs["Name"]:
    if x in jazz_tcs_list:
        # print("=======================\n" + x)
        different_flag = False
        # print(x)
        # print(word_tcs[word_tcs["Name"] == x]["Estimate"].tolist())
        # print(jazz_tcs[jazz_tcs["Name"] == x]["Estimate"].tolist())
        # print(word_tcs[word_tcs["Name"] == x]["Priority TC"].tolist())
        # print(jazz_tcs[jazz_tcs["Name"] == x]["Priority TC"].tolist())
        # print(word_tcs[word_tcs["Name"] == x]["Auto/Manual"].tolist())
        # print(jazz_tcs[jazz_tcs["Name"] == x]["Auto/Manual"].tolist())
        # print(x)
        TC = {  "Name" : x, 
                "Word Estimate" : word_tcs[word_tcs["Name"] == x]["Estimate"].tolist()[0],
                "Jazz Estimate" : jazz_tcs[jazz_tcs["Name"] == x]["Estimate"].tolist()[0],
                "Word Priority TC": word_tcs[word_tcs["Name"] == x]["Priority TC"].tolist()[0],
                "Jazz Priority TC" : jazz_tcs[jazz_tcs["Name"] == x]["Priority TC"].tolist()[0],
                "Word Auto/Manual": word_tcs[word_tcs["Name"] == x]["Auto/Manual"].tolist()[0],
                "Jazz Auto/Manual" : jazz_tcs[jazz_tcs["Name"] == x]["Auto/Manual"].tolist()[0],
                "Products Not In Word" : "",
                "Products Not On Jazz" : ""}
        # Compare estimate
        if TC["Word Estimate"] == TC ["Jazz Estimate"]:
            TC["Word Estimate"]=""
            TC ["Jazz Estimate"]=""
        else:
            different_flag = True

        # Compare priority
        if TC["Word Priority TC"] == TC ["Jazz Priority TC"]:
            TC["Word Priority TC"]=""
            TC ["Jazz Priority TC"]=""
        else: 
            different_flag = True

        # Compare Auto/Manual
        if TC["Word Auto/Manual"] == TC ["Jazz Auto/Manual"]:
            TC["Word Auto/Manual"]=""
            TC ["Jazz Auto/Manual"]=""
        else:
            different_flag = True

        # Compare Applicable products
        word_applicable_products = analyze_applicable_products_to_list(word_tcs[word_tcs["Name"] == x]["Applicable"].tolist()[0])
        jazz_applicable_products = analyze_applicable_products_to_list(jazz_tcs[jazz_tcs["Name"] == x]["Applicable"].tolist()[0])

        products_not_on_jazz_list = list(set(word_applicable_products).difference(set(jazz_applicable_products)))
        products_not_in_word_list = list(set(jazz_applicable_products).difference(set(word_applicable_products)))
        
        if products_not_in_word_list:
            different_flag = True
            for x in products_not_in_word_list:
                TC["Products Not In Word"] += (x + ",")
            TC["Products Not In Word"] = TC["Products Not In Word"][:-1]

        if products_not_on_jazz_list:
            different_flag = True
            for x in products_not_on_jazz_list:
                TC["Products Not On Jazz"] += (x + ",")
            TC["Products Not On Jazz"] = TC["Products Not On Jazz"][:-1]


        if different_flag:
            Difference_TCs.append(TC)
        else:
            TheSame_TCs.append({"Name" : x})

    
# Export Excel file
diff = pd.DataFrame(Difference_TCs)
same = pd.DataFrame(TheSame_TCs)

with pd.ExcelWriter("Analyzed.xlsx") as excelwriter:
    new_tcs.to_excel(excelwriter,sheet_name="Not On Jazz")
    diff.to_excel(excelwriter,sheet_name="Difference")
    same.to_excel(excelwriter,sheet_name="Same")


# a = ["PGS_B", "PGS_S", "PGS_C"]
# b = ["PGS_B", "PGS_S", "PGS_C"]

# c = list(set(a) ^ set(b))
# print(c)

# d = list(set(b).difference(set(a)))
# print(d)